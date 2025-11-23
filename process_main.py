import google.generativeai as genai
import os
import sys
import cv2
import warnings
import numpy as np
from pdf2image import convert_from_path
from PyPDF2 import PdfReader
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from PIL import Image, ImageDraw
import time

def process_mainflow(api_key, pdf_file_start_end_dict, options,min_area,paddling, additional_prompt, gemini_model):
    """
    pdf_file_start_end_dict: [
        {
            'name': 'file1.pdf',
            'start': 1,
            'end': 5
        },
        ...
    ]
    options: {
        'transcribe': True,
        'black_frame': False,
        'crop': True
    }
    """
    

    # ---------- Set Configuration ----------
    warnings.filterwarnings("ignore", category=RuntimeWarning)
    INPUT_FOLDER = "input"
    OUTPUT_FOLDER = "output"
    SETTING_FOLDER = "setting"
    config = read_conf(os.path.join(SETTING_FOLDER, "conf.txt"))
    poppler_path = config.get("poppler_path")
    generate_Img = options.get("crop")
    image_frame = options.get("black_frame")
    slowdown = options.get("slowdown")

    APIKey = api_key
    main_prompt = read_prompt(os.path.join(SETTING_FOLDER, "mainprompt.txt"),1)
    #additional_prompt = read_prompt("prompt.txt",0)
    final_prompt = main_prompt+additional_prompt

    all_jobs = []

    # ---------- Check paths ----------
    if not os.path.exists(INPUT_FOLDER):
            print("❌ input folder not found.")
            sys.stdout.flush()
            return False

    files = [f for f in os.listdir(INPUT_FOLDER) if os.path.isfile(os.path.join(INPUT_FOLDER, f))]
    if not files:
        print("❌ No files found in 'input' folder.")
        sys.stdout.flush()
        return False


    # ---handle image---
    if options.get("transcribe"):
        image_files = [f for f in files if f.lower().endswith((".png", ".jpg", ".jpeg"))]
        if image_files:
            all_jobs.append((
            [Image.open(os.path.join(INPUT_FOLDER, f)).convert("RGB") for f in sorted(image_files)],
            [1, len(image_files)],
            "picture"
        ))
        else:
            print("No image files detected in input.\n")
            sys.stdout.flush()

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    # ---------- Main process ----------
    pdf_files = [f for f in files if f.lower().endswith(".pdf")]

    for pdf_file in pdf_files:
        pdf_path = os.path.join(INPUT_FOLDER, pdf_file)

        if pdf_file in pdf_file_start_end_dict:
            start, end = pdf_file_start_end_dict[pdf_file]
        else:
            continue
        pages = list(range(start, end + 1))

        print("📄 Converting PDF pages to images...")
        sys.stdout.flush()

        try:
            imagesPDF = convert_from_path(
                pdf_path,
                dpi=200,
                fmt="png",
                poppler_path=poppler_path,
                first_page=pages[0] if pages else None,
                last_page=pages[-1] if pages else None,
            )
        except Exception as e:
            print("❌ Conversion failed:", e)
            sys.stdout.flush()
            return False
  
        if image_frame:
            for i, img in enumerate(imagesPDF):
                draw = ImageDraw.Draw(img)
                w, h = img.size
                thickness = 10
                draw.rectangle([0, 0, w-1, h-1], outline="black", width=thickness)
                imagesPDF[i] = img
        baseName = os.path.splitext(pdf_file)[0]
        all_jobs.append((imagesPDF, pages, baseName))
        print(f"✅ All selected pages in {pdf_file} have been successfully converted.\n")
        sys.stdout.flush()

        # ---------- If no API key, skip OCR ----------
        if not APIKey:
            print("No available API_key found in conf.txt. Skipping AI text extraction.")
            sys.stdout.flush()
            return False

    # ---------- Configure Gemini and perform OCR ----------
    genai.configure(api_key=APIKey)
    model = genai.GenerativeModel(gemini_model)
    print("模型:"+gemini_model)

    for singleimages, pages, fileName in all_jobs:
        fileName = os.path.splitext(fileName)[0]
        print(f"----📌Starting text extraction from {fileName}📌----\n")
        sys.stdout.flush()
        output_txt = os.path.join(OUTPUT_FOLDER, fileName+"content.txt")
        with open(output_txt, "w", encoding="utf-8") as f:
            for i, img in enumerate(singleimages, start=(pages[0] if pages else 1)):
                print(f"▶️Extracting text from {fileName} page {i}...")
                sys.stdout.flush()
                buffer = BytesIO()
                img.save(buffer, format="PNG")
                buffer.seek(0)
                try:
                    response = model.generate_content([
                        {"mime_type": "image/png", "data": buffer.getvalue()},
                        {"text": final_prompt}
                    ])
                    text = response.text.strip() if getattr(response, "text", None) else "[No text detected]"
                    if generate_Img:
                        extract_figures(img, OUTPUT_FOLDER, f"{fileName}_{i}",min_area, paddling)
                except Exception as e:
                    text = f"**********[Error extracting text: {e}]*********"
                f.write(f"{text}\n\n")
                print(f"☑️ {fileName}  Page {i} done.")
                sys.stdout.flush()
                if slowdown:
                    print("⚠️降速等待中...")
                    time.sleep(4)

        # ---------- Word File ----------

        txt_file = os.path.join(OUTPUT_FOLDER, fileName+"content.txt")
        if not os.path.exists(txt_file):
            print(f"❌ File not found: {txt_file}")
            sys.stdout.flush()
            return False

        with open(txt_file, "r", encoding="utf-8") as f:
            lines = f.readlines()

        doc = Document()

        for line in lines:
            line = line.rstrip("\n")
            if not line:
                doc.add_paragraph("")
                continue
            p = doc.add_paragraph()
            for char in line:
                run = p.add_run(char)
                if contains_chinese(char):
                    run.font.name = "標楷體"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
                else:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    if char.isalpha():  
                        run.italic = True
                run.font.size = Pt(12)

        output_file = os.path.join(OUTPUT_FOLDER, "(未校稿)"+fileName+".docx")
        doc.save(output_file)
        print(f"✅ Word file saved: {output_file}\n")
        sys.stdout.flush()
        os.remove(output_txt)
    os.startfile(OUTPUT_FOLDER)
    return True

def extract_figures(image, output_dir, base_name, mina, pad):
    os.makedirs(output_dir, exist_ok=True)

    img = np.array(image.convert("RGB"))
    img_gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

    _, thresh = cv2.threshold(img_gray, 240, 255, cv2.THRESH_BINARY_INV)

    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    boxes = []
    debug_img = img.copy()

    for i, cnt in enumerate(contours):
        x, y, w, h = cv2.boundingRect(cnt)
        if w * h < mina:
            continue

        boxes.append({"x": int(x), "y": int(y), "width": int(w), "height": int(h)})

        cv2.rectangle(debug_img, (x, y), (x + w, y + h), (255, 0, 0), 2)

        roi = img[y-pad:y+h+pad, x-pad:x+w+pad]
        out_path = os.path.join(output_dir, f"{base_name}_figure_{i+1}.png")
        Image.fromarray(roi).save(out_path)
        print(f"📸 已輸出輔助圖片: {out_path}")
        sys.stdout.flush()

# ---------- Read configuration ----------
def read_conf(file_path):
    config = {}
    if not os.path.exists(file_path):
        print(f"❌ {file_path} does not exist. Check setting folder again!")
        sys.stdout.flush()
        return False
    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            if ":" in line:
                key, value = line.split(":", 1)
                config[key.strip()] = value.strip().strip('"')
    return config
            
def read_prompt(file_path,mode):
    with open(file_path, "r", encoding="utf-8") as f:
        if not os.path.exists(file_path):
            print(f"{file_path} does not exist. Check again!")
            return False
        txt_cont = f.read().strip()
        if(mode==0 and txt_cont):
            return  "另外，我還需要:"+txt_cont
        else:
            return txt_cont

    
def contains_chinese(char):
    return '\u4e00' <= char <= '\u9fff'